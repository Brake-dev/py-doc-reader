from docx import Document
from structured_output import DocListSchema


def to_docx(doc_list: DocListSchema):
    document = Document("./docs/docx/template.docx")

    for topic in doc_list.topics:
        document.add_heading(topic.title, level=1)

        for question in topic.questions:
            document.add_paragraph(
                question, style="Bullet list 1"
            )  # Style is specific to the template file

    document.save("rfp_response.docx")
